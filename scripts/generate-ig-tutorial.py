from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

OUTPUT = '/home/z/my-project/download/Tutorial_Setup_Instagram_BG-Health.docx'

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.4

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0xFF, 0x4D, 0x00)  # BG-Health orange
    hs.font.bold = True
    if level == 1:
        hs.font.size = Pt(18)
        hs.paragraph_format.space_before = Pt(18)
    elif level == 2:
        hs.font.size = Pt(14)
        hs.paragraph_format.space_before = Pt(14)
    else:
        hs.font.size = Pt(12)
        hs.paragraph_format.space_before = Pt(10)

def add_step(num, title, content_lines):
    p = doc.add_paragraph()
    run = p.add_run(f'Step {num}: {title}')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xFF, 0x4D, 0x00)
    for line in content_lines:
        p = doc.add_paragraph(line, style='List Bullet')
        p.paragraph_format.space_after = Pt(2)

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    # light gray background
    shading = qn('w:shd')
    pPr = p._p.get_or_add_pPr()
    shd = pPr.makeelement(qn('w:shd'), {
        qn('w:fill'): 'F0F0F0',
        qn('w:val'): 'clear',
    })
    pPr.append(shd)

def add_note(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

def add_bold(label, text):
    p = doc.add_paragraph()
    r1 = p.add_run(label)
    r1.bold = True
    r2 = p.add_run(text)

# ═══════════════════════════════════════════════
#  CONTENT
# ═══════════════════════════════════════════════

doc.add_heading('Tutorial: Koneksi Instagram ke BG-Health v2', level=1)

p = doc.add_paragraph('Tutorial ini menjelaskan cara menghubungkan akun Instagram @Bagongnews ke halaman Home BG-Health v2 agar postingan Instagram tampil secara otomatis.')
p.paragraph_format.space_after = Pt(12)

# ── Prasyarat ──
doc.add_heading('Prasyarat', level=1)
doc.add_paragraph('Pastikan Anda memiliki:', style='List Bullet')
doc.add_paragraph('Akun Instagram @Bagongnews (sudah ada)', style='List Bullet')
doc.add_paragraph('Akun Facebook pribadi (untuk login ke Meta Developer)', style='List Bullet')
doc.add_paragraph('Akses ke Vercel project bg-health-v2 (untuk menambahkan Environment Variables)', style='List Bullet')

# ── Step 1 ──
doc.add_heading('Langkah 1: Buat Akun Meta Developer', level=1)
add_step(1, 'Buka Meta for Developers', [
    'Buka browser, kunjungi https://developers.facebook.com',
    'Login menggunakan akun Facebook Anda',
    'Jika pertama kali, setujui syarat dan ketentuan',
])

add_step(2, 'Buat Aplikasi Baru', [
    'Klik tombol “My Apps” di pojok kanan atas',
    'Klik “Create App”',
    'Pilih tipe: “Business”',
    'Isi nama app: “BG-Health Instagram Feed” (atau nama apapun)',
    'Klik “Create App”',
])

add_step(3, 'Konfigurasi Basic Settings', [
    'Di dashboard app, buka menu “Settings” → “Basic”',
    'Isi “App Contact Email” dengan email Anda',
    'Scroll ke bawah, klik “Add Platform” → pilih “Website”',
    'Isi URL: https://vercel.com (atau domain BG-Health Anda)',
    'Klik “Save Changes”',
])

# ── Step 2 ──
doc.add_heading('Langkah 2: Tambahkan Produk Instagram Graph API', level=1)
add_step(4, 'Tambahkan Produk', [
    'Di sidebar kiri dashboard app, klik “Add Product”',
    'Cari dan pilih “Instagram Graph API”',
    'Pilih mode: “Business” (jika akun IG adalah Business/Creator account)',
    'Klik “Set Up”',
])

add_step(5, 'Konfigurasi Instagram', [
    'Anda akan diarahkan ke halaman konfigurasi Instagram',
    'Pilih “Business Account” atau “Creator Account”',
    'Meta akan meminta Anda menghubungkan akun Instagram ke Facebook Page',
    'Jika @Bagongnews sudah terhubung ke Facebook Page, pilih Page tersebut',
    'Jika belum, Anda perlu membuat Facebook Page baru dan menghubungkannya ke akun Instagram @Bagongnews',
    'Klik “Continue” dan ikuti instruksi di layar',
])

add_note('Catatan: Akun Instagram @Bagongnews harus bertipe Business atau Creator Account. Jika masih Personal Account, ubah terlebih dahulu di Settings → Account → Switch to Professional Account di aplikasi Instagram.')

# ── Step 3 ──
doc.add_heading('Langkah 3: Buat Access Token', level=1)
add_step(6, 'Generate Token', [
    'Di dashboard app, buka menu “Tools & Plugins” di sidebar kiri',
    'Klik “Instagram Graph API Token Generator”',
    'Di bagian “Get User Access Token”:',
    'Pilih Facebook Page yang terhubung ke @Bagongnews',
    'Pilih Instagram Business Account: @Bagongnews',
])

add_step(7, 'Pilih Permission', [
    'Di dropdown “Permissions”, pastikan terpilih:',
    '   → instagram_basic (profil dasar)',
    '   → instagram_content_publish (opsional, untuk posting)',
    '   → pages_read_engagement (opsional, untuk statistik)',
    'Klik “Generate Token”',
    'Salin token yang muncul (panjang, berisi karakter alfanumerik)',
])

add_note('PENTING: Token ini bersifat sementara (biasanya 60 hari). Simpan baik-baik. Token yang kedaluwarsa akan menyebabkan feed tidak bisa dimuat. Untuk produksi, gunakan System User Token yang tidak pernah kedaluwarsa (lihat Catatan di bawah).')

add_step(8, 'Dapatkan Instagram User ID', [
    'Di halaman yang sama (Instagram Graph API Token Generator)',
    'Setelah token di-generate, akan muncul informasi akun',
    'Cari field “Instagram Business Account ID” atau “User ID”',
    'Salin angka ID tersebut (contoh: 17841400912345678)',
    'Atau, buka browser dan akses URL berikut (ganti TOKEN):',
])
add_code('https://graph.facebook.com/v21.0/me/accounts?fields=instagram_business_account{id,name}&access_token=TOKEN_ANDA')

# ── Step 4 ──
doc.add_heading('Langkah 4: Konfigurasi di Vercel', level=1)
add_step(9, 'Tambah Environment Variables', [
    'Buka https://vercel.com/dashboard',
    'Pilih project “bg-health-v2”',
    'Buka menu “Settings” → “Environment Variables”',
    'Tambahkan 2 variable baru:',
])

add_bold('INSTAGRAM_USER_ID = ', '17841400912345678 (ganti dengan ID Anda dari Step 8)')
add_bold('INSTAGRAM_ACCESS_TOKEN = ', 'IGQVJ...(ganti dengan token dari Step 7)')

add_note('Jangan gunakan prefix NEXT_PUBLIC_ karena ini adalah server-side secret yang tidak boleh terekspos ke browser.')

add_step(10, 'Redeploy', [
    'Setelah menambahkan environment variables, klik “Redeploy”',
    'Atau push commit baru ke GitHub untuk trigger deployment otomatis',
    'Tunggu deployment selesai (1-2 menit)',
])

# ── Step 5 ──
doc.add_heading('Langkah 5: Verifikasi', level=1)
add_step(11, 'Test API Endpoint', [
    'Buka browser dan akses:',
])
add_code('https://domain-anda.vercel.app/api/social-feed')
doc.add_paragraph('Response JSON seharusnya berisi data news dengan source: “instagram” dan thumbnail asli dari postingan Instagram.')

add_step(12, 'Cek Halaman Home', [
    'Buka halaman Home di BG-Health',
    'Bagian “News” seharusnya menampilkan postingan Instagram @Bagongnews',
    'Thumbnail dan caption harus sesuai dengan postingan asli',
    'Klik card seharusnya membuka postingan Instagram di tab baru',
])

# ── Troubleshooting ──
doc.add_heading('Troubleshooting', level=1)

add_bold('Error: “OAuthException - Session has expired”', '')
doc.add_paragraph('Token sudah kedaluwarsa. Ulangi Langkah 7 untuk generate token baru. Untuk solusi jangka panjang, gunakan System User Token (tidak pernah expired).')

add_bold('Error: “IG-API-1000 - User has not authorized application”', '')
doc.add_paragraph('Akun Instagram belum diotorisasi. Ulangi Langkah 5 untuk menghubungkan akun Instagram ke Facebook Page melalui app.')

add_bold('Error: “Unsupported get request”', '')
doc.add_paragraph('Pastikan Instagram User ID yang dimasukkan benar. Gunakan token generator tool di Meta Developer untuk mendapatkan ID yang tepat.')

add_bold('Feed masih menampilkan YouTube Podcast, bukan Instagram', '')
doc.add_paragraph('Pastikan kedua environment variables (INSTAGRAM_USER_ID dan INSTAGRAM_ACCESS_TOKEN) sudah ditambahkan di Vercel. Cek dengan membuka Settings → Environment Variables di Vercel dashboard.')

add_bold('Gambar tidak muncul (thumbnail kosong)”', '')
doc.add_paragraph('Beberapa postingan carousel atau video mungkin tidak memiliki thumbnail yang langsung bisa diakses. Ini normal untuk beberapa tipe media.')

# ── Token Permanen ──
doc.add_heading('Membuat Token Permanen (Opsional, untuk Produksi)', level=1)
doc.add_paragraph('Token yang didapat dari Token Generator bersifat sementara (60 hari). Untuk production, gunakan System User Token yang tidak pernah expired:')

add_step(13, 'Buat System User', [
    'Di Meta App Dashboard, buka “Users → System Users”',
    'Klik “Add System User”',
    'Isi nama: “BG-Health Feed Reader”',
    'Pilih role: “Instagram Content Reader” atau “Admin”',
    'Klik “Create System User”',
])

add_step(14, 'Generate System Token', [
    'Setelah System User dibuat, klik “Generate New Token”',
    'Pilih permissions: instagram_basic',
    'Klik “Generate Token”',
    'SALIN TOKEN INI - token ini tidak kedaluwarsa',
    'Gunakan token ini sebagai INSTAGRAM_ACCESS_TOKEN di Vercel',
])

add_note('System User Token tidak memiliki batas waktu kedaluwarsa. Namun, jika Anda mengubah app secret atau menghapus System User, token akan menjadi tidak valid.')

# ── Ringkasan ──
doc.add_heading('Ringkasan File yang Perlu Di-update', level=1)
doc.add_paragraph('File berikut sudah di-update dan siap deploy:')
doc.add_paragraph('src/app/api/social-feed/route.ts — API endpoint dengan dukungan Instagram Graph API', style='List Bullet')
doc.add_paragraph('.env.example — Daftar environment variables yang diperlukan', style='List Bullet')
doc.add_paragraph('Tutorial_Setup_Instagram_BG-Health.docx — Dokumen ini', style='List Bullet')

doc.add_paragraph('Environment Variables yang perlu ditambahkan di Vercel:')
add_code('INSTAGRAM_USER_ID=17841400912345678\nINSTAGRAM_ACCESS_TOKEN=IGQVJ...')

doc.add_paragraph('Jika env vars tidak diisi, bagian News akan otomatis menampilkan video YouTube Podcast dari channel Bagong News sebagai fallback.')

# ── Save ──
doc.save(OUTPUT)
print(f'Tutorial saved to {OUTPUT}')
print(f'File size: {os.path.getsize(OUTPUT)} bytes')
